"""Build the one-page Fiji/ImageJ cross-check methods appendix."""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "report" / "Fiji_ImageJ_复核方法附录.docx"
BLUE = "2E74B5"
INK = "1F4D78"
LIGHT = "E8EEF5"
GRAY = "F4F6F9"


def set_font(run, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = OxmlElement("w:shd")
    node.set(qn("w:fill"), fill)
    tc_pr.append(node)


def set_cell_width(cell, inches: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def fixed_table(table, widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            margins = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcMar")
            if margins is None:
                margins = OxmlElement("w:tcMar")
                cell._tc.get_or_add_tcPr().append(margins)
            for side in ("top", "start", "bottom", "end"):
                item = margins.find(qn(f"w:{side}"))
                if item is None:
                    item = OxmlElement(f"w:{side}")
                    margins.append(item)
                item.set(qn("w:w"), "65" if side in {"top", "bottom"} else "100")
                item.set(qn("w:type"), "dxa")


def add_text(cell, text: str, *, bold: bool = False, size: float = 8.2, color: str = "000000") -> None:
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def paragraph(doc: Document, text: str, *, size: float = 9.2, bold: bool = False, color: str = "000000", after: float = 3.5) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)


def heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    set_font(run, size=10.4, bold=True, color=BLUE)


def main() -> None:
    summary = pd.read_csv(ROOT / "outputs" / "fiji_review_summary.csv").set_index("metric")
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.28)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.2)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("POROUS COATING SEM MORPHOMETRY  |  METHODS APPENDIX")
    set_font(run, 7.3, bold=True, color="667085")

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(1)
    run = title.add_run("Fiji/ImageJ 复核：SEM 形貌定量方法附录")
    set_font(run, 16, bold=True, color=INK)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(5)
    run = sub.add_run("冻结的 12 张分层随机 SEM 图像（seed = 20260830）；与项目 Python 基线逐图对照")
    set_font(run, 8.4, color="4B5563")

    callout = doc.add_table(rows=1, cols=1)
    fixed_table(callout, [7.4])
    shade(callout.cell(0, 0), GRAY)
    add_text(callout.cell(0, 0), "结论  孔隙面积分数可互换地复核；等效直径与圆度保留高度一致的排序；对象数对 Watershed 和小连通域规则敏感，不应在两软件间直接混用。", bold=False, size=8.65, color=INK)

    heading(doc, "1  冻结抽样与等价处理")
    paragraph(doc, "从 28 张已校准 SEM/SEM-like 图像按图像角色分层随机抽取 12 张：coating 7、particle 4、reference 1。导出的是 Python 主流程同一裁剪 ROI 与 1–99% 对比度归一化版本；TIFF 保留 8-bit 灰度，PNG 复用 manual_calibration.csv。Fiji 对每图执行 Analyze › Set Scale（1 px = manifest 中 pixel_size_um），不使用 Global scale。", after=2.5)
    paragraph(doc, "Fiji 宏：Gaussian Blur σ=1 px → Otsu（亮对象）→ Convert to Mask → binary Open → Watershed → Analyze Particles。最小面积使用与 Python 相同的图像自适应阈值（max[12 px, 1.5×10⁻⁵×图像像素数]）并换算为 µm²。孔隙率在 Watershed 前由前景像素分数计算；直径、圆度和对象数在 Watershed 后测量。", after=2.5)

    heading(doc, "2  逐图对照结果（n = 12；Fiji − Python）")
    table = doc.add_table(rows=1, cols=5)
    fixed_table(table, [1.62, 1.25, 1.25, 1.55, 1.73])
    headers = ["指标", "Spearman ρ", "中位 |Δ|", "95% bootstrap CI", "解释"]
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, LIGHT)
        add_text(cell, text, bold=True, size=8.0, color=INK)
    rows = [
        ("孔隙面积分数", "pore_area_fraction", "面积分数"),
        ("等效直径 (µm)", "eq_diameter_median_um", "µm"),
        ("圆度", "circularity_median", "无量纲"),
        ("对象数", "object_count", "个/图"),
    ]
    interpretations = ["一致，可作面积分数复核", "排序一致；尺度有系统差", "排序一致；边界估计差异小", "不一致；分水岭种子敏感"]
    for (label, key, unit), explain in zip(rows, interpretations):
        s = summary.loc[key]
        cells = table.add_row().cells
        values = [
            label,
            f"{s.spearman_rho:.2f}",
            f"{s.median_absolute_difference:.4g} {unit}",
            f"[{s.median_absolute_difference_bootstrap_ci_low:.3g}, {s.median_absolute_difference_bootstrap_ci_high:.3g}]",
            explain,
        ]
        for cell, text in zip(cells, values):
            add_text(cell, text, size=7.75)

    heading(doc, "3  一致的条件与失败边界")
    paragraph(doc, "一致：同一 ROI、同一像素标定、对比度稳定且目标相对背景为亮相时，Otsu 得到的面积分数几乎等价（ρ=1.00，中位 |Δ|=0.00213）。粒径与圆度的 ρ=0.99，适合比较相对排序或趋势。", after=1.8)
    paragraph(doc, "失败：对象相互接触、边界弱、强充电/阴影、倍率导致颗粒跨越最小面积门槛，或连通域在 Watershed 中被不同种子切分时，对象数会偏离（ρ=0.51；Fiji 中位数高 217 个/图）。此时应报告掩膜、保持单一软件的计数规则，并将分割误差作为方法限制，而非解释为真实制备差异。", after=2)
    source = doc.add_paragraph()
    source.paragraph_format.space_before = Pt(2)
    source.paragraph_format.space_after = Pt(0)
    source.paragraph_format.line_spacing = 1.0
    run = source.add_run("可复现文件：fiji/sem_morphometry_review.ijm；scripts/build_fiji_review_package.py；scripts/compare_fiji_python.py；outputs/fiji_review_comparison.csv。ImageJ 官方文档：imagej.net/tutorials/batch-processing-with-ij-macro。")
    set_font(run, 7.1, color="667085")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Fiji/ImageJ cross-check  |  v0.2.0 extension")
    set_font(run, 7.2, color="667085")
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
