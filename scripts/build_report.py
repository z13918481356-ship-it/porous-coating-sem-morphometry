from __future__ import annotations

import argparse
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT = Path(__file__).resolve().parents[1]
NAVY = RGBColor(22, 50, 79)
BLUE = RGBColor(47, 102, 144)
GRAY = RGBColor(92, 101, 112)
RED = RGBColor(181, 82, 69)


def set_font(run, size=None, bold=None, color=None, italic=None):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.bold = bold
    if color is not None: run.font.color.rgb = color
    if italic is not None: run.italic = italic


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal.font.size = Pt(11); normal.font.color.rgb = RGBColor(25, 25, 25)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_before = Pt(0); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for name, size, before, after, color_value in [
        ("Heading 1", 16, 16, 8, BLUE), ("Heading 2", 13, 12, 6, BLUE), ("Heading 3", 12, 8, 4, NAVY)
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"; style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = color_value
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 "); set_font(run, 8.5, color=GRAY)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r"); t = OxmlElement("w:t"); t.text = "1"; r.append(t); fld.append(r); paragraph._p.append(fld)
    run = paragraph.add_run(" / 2 页"); set_font(run, 8.5, color=GRAY)


def add_para(doc, text, *, size=11, bold=False, color=None, italic=False, after=6, align=None):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.10
    if align is not None: p.alignment = align
    run = p.add_run(text); set_font(run, size, bold, color, italic)
    return p


def add_callout(doc, label, text, color=BLUE):
    p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(.18); p.paragraph_format.right_indent = Inches(.12)
    p.paragraph_format.space_before = Pt(3); p.paragraph_format.space_after = Pt(8); p.paragraph_format.line_spacing = 1.08
    pPr = p._p.get_or_add_pPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F4F7"); pPr.append(shd)
    r = p.add_run(label + "  "); set_font(r, 10.5, True, color)
    r = p.add_run(text); set_font(r, 10.5)


def add_picture(doc, path: Path, width: float, caption: str):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("descr", caption)
    picture._inline.docPr.set("title", path.stem)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(5)
    r = cap.add_run(caption); set_font(r, 8.5, color=GRAY, italic=True)


def build(output_root: Path, report_root: Path):
    summary = json.loads((output_root / "run_summary.json").read_text(encoding="utf-8"))
    models = pd.read_csv(output_root / "small_data_model_results.csv")
    associations = pd.read_csv(output_root / "bootstrap_intervals.csv")
    morph = pd.read_csv(output_root / "morphometry_features.csv")
    properties = pd.read_csv(output_root / "properties_clean.csv")
    unet_summary = pd.read_csv(output_root / "fibsem_unet_test_summary.csv")
    unet_bootstrap = pd.read_csv(output_root / "fibsem_unet_paired_bootstrap.csv").iloc[0]
    rf_ca = models[(models.target == "contact_angle_deg") & (models.model == "Random Forest")].iloc[0]
    baseline_ca = models[(models.target == "contact_angle_deg") & (models.model == "Training-mean baseline")].iloc[0]
    top = associations[associations.target == "contact_angle_deg"].copy()
    top["abs_rho"] = top.spearman_rho.abs(); top = top.sort_values("abs_rho", ascending=False).iloc[0]
    sensitivity = (morph.solid_fraction_sensitivity_max - morph.solid_fraction_sensitivity_min).median()
    unet_all = unet_summary[(unet_summary.method == "unet") & (unet_summary.stratum == "ALL")].iloc[0]
    unet_hpc22 = unet_summary[(unet_summary.method == "unet") & (unet_summary.stratum == "HPC22")].iloc[0]
    otsu_all = unet_summary[(unet_summary.method == "otsu") & (unet_summary.stratum == "ALL")].iloc[0]

    doc = Document(); section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(.492)
    configure_styles(doc)

    header = section.header.paragraphs[0]; header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("RESEARCH BRIEF  ·  SEM MORPHOMETRY"); set_font(r, 8.5, True, GRAY)
    add_page_field(section.footer.paragraphs[0])

    add_para(doc, "阶段 1 / 实证分析", size=9.5, bold=True, color=BLUE, after=2)
    add_para(doc, "多孔功能涂层的定量 SEM 形貌与结构—性能映射", size=22, bold=True, color=NAVY, after=4)
    add_para(doc, "严格限定为形貌—润湿/耐久性关联；不包含发射率或其他光学性质预测", size=11.5, color=GRAY, after=10)
    add_para(doc, "主数据：Zenodo 16054027  |  外部基准：5905496 / 4317170  |  许可：CC BY 4.0", size=9.2, color=GRAY, after=10)

    add_callout(doc, "结论", f"可从多尺度 SEM 中提取颗粒/孔隙描述符并与润湿、磨损状态做探索性关联；但质量控制后只有 {summary['matched_condition_groups']} 个独立匹配条件，置换检验不支持可部署的性质预测。独立 FIB-SEM 基准证明监督分割优于全局阈值，但仍未达到预注册可靠性门槛。")
    doc.add_heading("问题、数据与边界", level=1)
    add_para(doc, f"公开档案含 {summary['images_total']} 幅 SEM/SEM 类图像，其中 {summary['coating_images']} 幅为涂层图像；全部完成物理尺度校准。清理 331 个 Excel/资源文件后保留 {summary['unique_property_workbooks']} 个唯一有效性质工作簿。图像与性质仅在制备条件、测试类型和循环数可无歧义对应时合并，共 {summary['matched_modeling_images']} 幅图、{summary['matched_condition_groups']} 个独立条件。")
    add_para(doc, "分析单元已拆分为条件、试片、SEM 视野和性质工作簿四层；源档案未提供试片 ID，故试片层保持为空。同一条件的高/低倍率图像不跨数据划分，分割失败视野不进入模型。粒子亮区阈值可近似表征固相覆盖，但暗区同时包含真实孔隙、阴影与起伏，不能解释为三维孔隙率。", after=7)
    add_picture(doc, output_root / "figures" / "figure_2_segmentation_comparison.png", 6.35,
                "图 1  主涂层 SEM 的 Otsu、形态学与 Watershed 输出；主数据无专家掩膜，因此不报告涂层域 U-Net 准确率。")

    doc.add_page_break()
    doc.add_heading("结果与可靠性", level=1)
    add_callout(doc, "量化摘要", f"中位阈值敏感跨度为 {sensitivity:.3f}（固相面积分数）；{summary['failure_cases']} 幅图触发失败标记并被排除。接触角仅 {int(rf_ca.n_conditions)} 个条件：Random Forest 留一 MAE={rf_ca.mae:.2f}°，均值基线={baseline_ca.mae:.2f}°，置换 p={rf_ca.permutation_p_one_sided:.3f}。", color=NAVY)
    add_picture(doc, output_root / "figures" / "figure_4_wetting_durability.png", 5.25,
                "图 2  PDMS/PUR、烧结状态与磨损方式下的接触角和滞后角轨迹；每条线是条件级汇总。")
    doc.add_heading("形貌—性质关联与模型", level=2)
    add_para(doc, f"在 {summary['matched_condition_groups']} 个匹配条件中，接触角与 {top.feature.replace('_', ' ')} 的 Spearman ρ={top.spearman_rho:.2f}，组 bootstrap 95% CI [{top.ci_low:.2f}, {top.ci_high:.2f}]，只能生成假设。Random Forest 略低于均值基线，但置换 p={rf_ca.permutation_p_one_sided:.3f}；滞后角与滚动角各仅 4 条件，未拟合。", size=10.2, after=4)
    doc.add_heading("独立分割基准", level=2)
    add_callout(doc, "冻结测试", f"官方 180/60/60 划分上，U-Net 测试 Dice 中位数由 Otsu 的 {otsu_all.dice_median:.3f} 提升至 {unet_all.dice_median:.3f}；配对增益 {unet_bootstrap.median_dice_gain:.3f}，95% CI [{unet_bootstrap.median_dice_gain_ci_low:.3f}, {unet_bootstrap.median_dice_gain_ci_high:.3f}]，胜出 {int(unet_bootstrap.unet_wins)}/60。总体仍低于 0.85 门槛，且 HPC22={unet_hpc22.dice_median:.3f}。", color=NAVY)
    add_picture(doc, output_root / "external_validation" / "fibsem_unet_training_and_comparison.png", 4.8,
                "图 3  验证集选模与官方测试集配对比较；测试集仅在第 26 轮检查点冻结后评估一次。")
    doc.add_heading("判断与下一步", level=2)
    add_para(doc, "涂层域 24 个双人标注 patch 仍待专家完成；FIB-SEM 与 PAAO 只作分割压力测试，不进入润湿模型。若继续改模型，须新增未触碰评估集，不能把官方测试集变成开发集。光学模拟仍需可靠光学常数与显式几何假设。", size=9.6, after=3)
    add_para(doc, "来源：Sultan et al. (2025), Zenodo 16054027；PAAO 5905496；FIB-SEM 4317170。代码、数据卡、5 张主分析图、外部诊断图及失败案例位于同一仓库。", size=8.2, color=GRAY, after=0)

    report_root.mkdir(parents=True, exist_ok=True)
    out = report_root / "SEM形貌_润湿耐久性关联_两页报告.docx"
    doc.save(out)
    print(out)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--output-root", type=Path, default=PROJECT / "outputs")
    parser.add_argument("--report-root", type=Path, default=PROJECT / "report")
    args = parser.parse_args(); build(args.output_root, args.report_root)
